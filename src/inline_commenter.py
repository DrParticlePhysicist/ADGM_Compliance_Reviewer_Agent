from docx import Document
from pathlib import Path
from docx.shared import RGBColor


def insert_comments(docx_path: Path, issues: list, output_path: Path):
    """
    Insert pseudo-inline comments next to the relevant text in the original content.

    Args:
        docx_path (Path): Path to the original docx file
        issues (list): List of dicts with at least "issue" key, optionally "severity", "suggestion", "match_text"
        output_path (Path): Path to save the reviewed docx file
    """
    doc = Document(docx_path)

    for issue in issues:
        severity = issue.get('severity', 'N/A')
        short_issue = issue.get('issue', 'No issue provided')
        suggestion = issue.get('suggestion', issue.get('explanation', 'No suggestion provided'))
        match_text = issue.get('match_text', short_issue)

        comment_text = f"  [COMMENT - Severity: {severity} | {short_issue} | Suggestion: {suggestion}]"

        inserted = False
        for para in doc.paragraphs:
            if match_text and match_text in para.text:
                run = para.add_run(comment_text)
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for visibility
                inserted = True
                break

        # If not found, add comment at end of document as fallback
        if not inserted:
            p = doc.add_paragraph(comment_text)
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    doc.save(output_path)
